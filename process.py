import shutil
import datetime
from docx import Document
import glob
from docx2pdf import convert
from pdf2image import convert_from_path
import csv
import sys
import os
from tkinter import messagebox
import tkinter as tk

root = tk.Tk()
root.attributes('-topmost', True)
root.withdraw()
root.lift()
root.focus_force()


def get_dir_path():
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(__file__)

    return base_path


CSV_DIR = get_dir_path() + "/csv"
TEMPLATE_DOC = get_dir_path() + "/document/template.docx"
DOC_DIR = get_dir_path() + "/document/temp"
PDF_DIR = get_dir_path() + "/pdf"
IMAGES_DIR = get_dir_path() + "/images"
POPPLER_DIR = get_dir_path() + "/poppler-0.67.0/bin"


def get_csv_data_list():
    """
    csvファイル一覧を取得し、値を配列で返す
    """

    try:
        csvFilePathList = glob.glob(CSV_DIR + "/*.csv")

        csvDataList = []
        for csvFilePath in csvFilePathList:
            csvFile = open(csvFilePath, "r",
                           encoding="ms932", errors="", newline="")
            f = csv.reader(csvFile, delimiter=",", doublequote=True,
                           lineterminator="\r\n", quotechar='"', skipinitialspace=True)
            for row in f:
                csvDataList.append(row)

        return csvDataList
    except Exception as e:
        messagebox.showerror("処理失敗", "csvデータ取得処理失敗")
        raise (e)


def create_document():
    """
    wordファイルを作成
    """

    try:
        now = datetime.datetime.now(datetime.timezone(
            datetime.timedelta(hours=9)))
        nowStr = now.strftime('%Y%m%d%H%M%S')

        csvDataList = get_csv_data_list()

        for row in csvDataList:
            if len(row) == 3:
                new_doc_path = DOC_DIR + "/" + row[2] + "_" + nowStr + ".docx"
                shutil.copyfile(TEMPLATE_DOC, new_doc_path)

                doc = Document(new_doc_path)

                # 行単位の置換
                for paragraph in doc.paragraphs:
                    if paragraph.text == "${date}":
                        paragraph.text = row[0]
                    if paragraph.text == "依頼債務者　${requesterName}":
                        paragraph.text = "依頼債務者　" + row[1]
                    if paragraph.text == "　　　　${merchantName}　　御　中":
                        paragraph.text = "　　　　" + row[2] + "　　御　中"

                paragraphs = (paragraph
                              for table in doc.tables
                              for row in table.rows
                              for cell in row.cells
                              for paragraph in cell.paragraphs)

                # テーブルのセル単位の置換
                for paragraph in paragraphs:
                    if paragraph.text == "${date}":
                        paragraph.text = row[0]
                    if paragraph.text == "依頼債務者　${requesterName}":
                        paragraph.text = "依頼債務者　" + row[1]
                    if paragraph.text == "　　　　${merchantName}　　御　中":
                        paragraph.text = "　　　　" + row[2] + "　　御　中"

                doc.save(new_doc_path)
    except Exception as e:
        messagebox.showerror("処理失敗", "wordファイル作成処理処理失敗")
        raise (e)


def create_pdf():
    """
    wordファイルをpdfファイルに変換
    """

    try:
        docFilePathList = glob.glob(DOC_DIR + "/*.docx")

        for docFilePath in docFilePathList:

            fileNameSplit = docFilePath.split("\\")
            fileName = fileNameSplit[len(fileNameSplit) - 1].split(".")[0]

            outputFile = PDF_DIR + "/" + fileName + ".pdf"
            file = open(outputFile, "w")
            file.close()

            convert(docFilePath, outputFile)
    except Exception as e:
        messagebox.showerror("処理失敗", "pdfファイル作成処理処理失敗")
        raise (e)


def create_image():
    """
    pdfファイルをjpegファイルに変換
    """

    try:
        pdfFilePathList = glob.glob(PDF_DIR + "/*.pdf")

        for pdfFilePath in pdfFilePathList:
            fileNameSplit = pdfFilePath.split("\\")
            fileName = fileNameSplit[len(fileNameSplit) - 1].split(".")[0]

            pages = convert_from_path(pdfFilePath, poppler_path=POPPLER_DIR)
            for page in pages:
                page.save(IMAGES_DIR + "/" + fileName + '.jpg', 'JPEG')
    except Exception as e:
        messagebox.showerror("処理失敗", "画像ファイル作成処理処理失敗")
        raise (e)


try:
    shutil.rmtree(DOC_DIR)
    os.mkdir(DOC_DIR)
    shutil.rmtree(PDF_DIR)
    os.mkdir(PDF_DIR)
    shutil.rmtree(IMAGES_DIR)
    os.mkdir(IMAGES_DIR)

    try:
        create_document()
        create_pdf()
        create_image()

        messagebox.showinfo("処理成功", "処理完了!!")
    except Exception as e:
        print(e)
except Exception as e:
    print(e)
    messagebox.showerror("処理失敗", "ファイル削除失敗")
